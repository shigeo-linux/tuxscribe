import os
import re


def _format_reference(r):
    title = r.get('cite_title') or r.get('filename', '')
    author = r.get('cite_author') or ''
    year = r.get('cite_year') or ''
    city = r.get('cite_city') or ''
    publisher = r.get('cite_publisher') or ''
    parts = []
    if author:
        parts.append(author + '.')
    if title:
        parts.append(title + '.')
    location = ', '.join(filter(None, [publisher, city]))
    if location:
        parts.append(location + (',' if year else '.'))
    if year:
        parts.append(year + '.')
    return ' '.join(parts) if parts else title


def _chapter_heading(ch):
    num = ch['chapter_number']
    title = (ch['title'] or '').strip()
    if title:
        return f"Chapter {num}: {title}"
    return f"Chapter {num}"


def _text_to_html(text):
    paras = text.split('\n\n')
    parts = []
    for p in paras:
        p = p.strip()
        if p:
            p = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            p = p.replace('\n', '<br/>')
            parts.append(f'<p>{p}</p>')
    return '\n'.join(parts)


def export_epub(chapters, title, author, output_path, sources=None):
    try:
        from ebooklib import epub
    except ImportError:
        raise RuntimeError("ebooklib not installed.\nRun: pip3 install ebooklib")

    book = epub.EpubBook()
    book.set_identifier(f'tuxscribe-{re.sub(r"[^a-z0-9]", "-", title.lower())}')
    book.set_title(title)
    book.set_language('en')
    if author:
        book.add_author(author)

    css_content = b"""
body { font-family: Georgia, serif; font-size: 1em; line-height: 1.7;
       margin: 5% 10%; color: #111; }
h1 { font-size: 1.4em; margin: 0 0 1.5em 0; font-weight: normal;
     border-bottom: 1px solid #ccc; padding-bottom: 0.4em; }
p { text-indent: 1.5em; margin: 0 0 0.3em 0; }
p:first-of-type, p.noindent { text-indent: 0; }
"""
    css_item = epub.EpubItem(
        uid='style', file_name='style/main.css',
        media_type='text/css', content=css_content
    )
    book.add_item(css_item)

    epub_chapters = []
    for i, ch in enumerate(chapters):
        content = ch['content'] or ''
        if not content.strip():
            continue
        ch_title = _chapter_heading(ch)
        ec = epub.EpubHtml(
            title=ch_title,
            file_name=f'chap_{i+1:03d}.xhtml',
            lang='en',
        )
        ec.content = f'<h1>{ch_title}</h1>\n{_text_to_html(content)}'
        ec.add_item(css_item)
        book.add_item(ec)
        epub_chapters.append(ec)

    if sources:
        ref_html = '\n'.join(
            f'<p class="noindent">{_format_reference(r)}</p>' for r in sources
        )
        ref_page = epub.EpubHtml(title='References', file_name='references.xhtml', lang='en')
        ref_page.content = f'<h1>References</h1>\n{ref_html}'
        ref_page.add_item(css_item)
        book.add_item(ref_page)
        epub_chapters.append(ref_page)

    book.toc = tuple(epub_chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ['nav'] + epub_chapters
    epub.write_epub(output_path, book)


def export_docx(chapters, title, author, output_path, sources=None):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed.\nRun: pip3 install python-docx")

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True

    if author:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f'by {author}')
        r2.font.size = Pt(14)

    doc.add_page_break()

    for ch in chapters:
        content = ch['content'] or ''
        if not content.strip():
            continue
        ch_title = _chapter_heading(ch)
        doc.add_heading(ch_title, level=1)

        first = True
        for para_text in content.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(0)
            if not first:
                p.paragraph_format.first_line_indent = Inches(0.4)
            first = False

        doc.add_page_break()

    if sources:
        doc.add_heading('References', level=1)
        for r in sources:
            p = doc.add_paragraph(_format_reference(r))
            p.paragraph_format.space_after = Pt(4)

    doc.save(output_path)


def export_pdf(chapters, title, author, output_path, sources=None):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak
        )
    except ImportError:
        raise RuntimeError("reportlab not installed.\nRun: pip3 install reportlab")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=1.25*inch, rightMargin=1.25*inch,
        topMargin=1.2*inch, bottomMargin=1.2*inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'BTitle', parent=styles['Title'],
        fontSize=26, spaceAfter=16, alignment=TA_CENTER,
    )
    author_style = ParagraphStyle(
        'BAuthor', parent=styles['Normal'],
        fontSize=13, spaceAfter=8, alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        'BHeading', parent=styles['Heading1'],
        fontSize=15, spaceBefore=0, spaceAfter=18,
    )
    body_style = ParagraphStyle(
        'BBody', parent=styles['Normal'],
        fontSize=11, leading=17,
        alignment=TA_JUSTIFY,
        firstLineIndent=28, spaceAfter=2,
    )
    body_first_style = ParagraphStyle(
        'BBodyFirst', parent=body_style, firstLineIndent=0,
    )

    def esc(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    story = [
        Spacer(1, 1.8*inch),
        Paragraph(esc(title), title_style),
    ]
    if author:
        story.append(Paragraph(esc(f'by {author}'), author_style))
    story.append(PageBreak())

    for ch in chapters:
        content = ch['content'] or ''
        if not content.strip():
            continue
        ch_title = _chapter_heading(ch)
        story.append(Paragraph(esc(ch_title), heading_style))

        first = True
        for para_text in content.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            story.append(Paragraph(esc(para_text),
                                   body_first_style if first else body_style))
            first = False

        story.append(PageBreak())

    if sources:
        story.append(Paragraph('References', heading_style))
        for r in sources:
            story.append(Paragraph(esc(_format_reference(r)), body_first_style))
            story.append(Spacer(1, 4))

    doc.build(story)
